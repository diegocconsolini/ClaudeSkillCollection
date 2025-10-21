#!/usr/bin/env python3
"""
Convert Guatemaltek markdown policy files to professionally formatted Word documents
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
from datetime import datetime

def add_page_number(section):
    """Add page numbers to footer"""
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._element.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar2)

    footer_para.runs[0].font.size = Pt(10)
    footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

def setup_styles(doc):
    """Configure professional document styles"""
    styles = doc.styles

    # Title style
    if 'CustomTitle' not in [s.name for s in styles]:
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(26)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        title_style.paragraph_format.space_after = Pt(12)
        title_style.paragraph_format.space_before = Pt(0)

    # Heading 1
    heading1 = styles['Heading 1']
    heading1.font.name = 'Calibri'
    heading1.font.size = Pt(18)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 51, 102)
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(6)
    heading1.paragraph_format.keep_with_next = True

    # Heading 2
    heading2 = styles['Heading 2']
    heading2.font.name = 'Calibri'
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 76, 153)
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(4)
    heading2.paragraph_format.keep_with_next = True

    # Heading 3
    heading3 = styles['Heading 3']
    heading3.font.name = 'Calibri'
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.font.color.rgb = RGBColor(0, 102, 204)
    heading3.paragraph_format.space_before = Pt(8)
    heading3.paragraph_format.space_after = Pt(3)
    heading3.paragraph_format.keep_with_next = True

    # Normal text
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

def add_header_footer(doc, company="Guatemaltek", policy_name=""):
    """Add professional header and footer"""
    section = doc.sections[0]

    # Header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Company name in header
    run = header_para.add_run(company)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(64, 64, 64)

    # Policy name on right side
    if policy_name:
        tab_stops = header_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.0), WD_ALIGN_PARAGRAPH.RIGHT)
        header_para.add_run('\t')
        run = header_para.add_run(policy_name)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(64, 64, 64)

    # Add horizontal line
    header_para.paragraph_format.border_bottom = True

    # Footer with page numbers
    add_page_number(section)

def parse_metadata(content):
    """Extract metadata from document header"""
    metadata = {}
    lines = content.split('\n')

    for line in lines[:15]:
        if '**Company:**' in line:
            metadata['company'] = re.search(r'\*\*Company:\*\*\s*(.+?)(?:\s*\||\s*$)', line)
            if metadata['company']:
                metadata['company'] = metadata['company'].group(1).strip()
        if '**Version:**' in line or '**V:**' in line:
            metadata['version'] = re.search(r'\*\*(?:Version|V):\*\*\s*(.+?)(?:\s*\||\s*$)', line)
            if metadata['version']:
                metadata['version'] = metadata['version'].group(1).strip()
        if '**Effective Date:**' in line or '**Effective:**' in line:
            metadata['effective'] = re.search(r'\*\*Effective(?:\s+Date)?:\*\*\s*(.+?)(?:\s*\||\s*$)', line)
            if metadata['effective']:
                metadata['effective'] = metadata['effective'].group(1).strip()
        if '**Review Schedule:**' in line or '**Review:**' in line:
            metadata['review'] = re.search(r'\*\*Review(?:\s+Schedule)?:\*\*\s*(.+?)(?:\s*\||\s*$)', line)
            if metadata['review']:
                metadata['review'] = metadata['review'].group(1).strip()
        if '**Responsible Officer:**' in line or '**Officer:**' in line:
            metadata['officer'] = re.search(r'\*\*(?:Responsible\s+)?Officer:\*\*\s*(.+?)(?:\s*\||\s*$)', line)
            if metadata['officer']:
                metadata['officer'] = metadata['officer'].group(1).strip()

    return metadata

def add_metadata_table(doc, metadata):
    """Add formatted metadata table"""
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Light Grid Accent 1'

    fields = [
        ('Company', metadata.get('company', 'Guatemaltek')),
        ('Version', metadata.get('version', '1.0')),
        ('Effective Date', metadata.get('effective', '2026-01-01')),
        ('Review Schedule', metadata.get('review', 'Annually')),
        ('Responsible Officer', metadata.get('officer', 'CISO'))
    ]

    for field, value in fields:
        row = table.add_row()
        row.cells[0].text = field
        row.cells[1].text = value

        # Bold field names
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()  # Space after table

def convert_markdown_to_docx(md_file, output_file):
    """Convert markdown file to formatted Word document"""

    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create new document
    doc = Document()

    # Setup styles
    setup_styles(doc)

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Parse metadata
    metadata = parse_metadata(content)

    # Extract policy name from first line (title)
    lines = content.split('\n')
    policy_name = lines[0].replace('#', '').strip() if lines else "Policy"

    # Add header and footer
    add_header_footer(doc, metadata.get('company', 'Guatemaltek'), policy_name)

    # Process content
    in_code_block = False
    in_list = False
    list_level = 0
    skip_metadata = True

    for line in lines:
        # Skip initial metadata section
        if skip_metadata and line.strip() == '---':
            skip_metadata = False
            continue

        if skip_metadata and (line.startswith('**') or not line.strip()):
            continue

        if line.startswith('---'):
            continue

        # Title (H1)
        if line.startswith('# '):
            title = line.replace('#', '').strip()
            para = doc.add_paragraph(title, style='CustomTitle')

            # Add metadata table after title
            add_metadata_table(doc, metadata)
            continue

        # Headings
        if line.startswith('## '):
            heading = line.replace('##', '').strip()
            doc.add_paragraph(heading, style='Heading 1')
            in_list = False
            continue

        if line.startswith('### '):
            heading = line.replace('###', '').strip()
            doc.add_paragraph(heading, style='Heading 2')
            in_list = False
            continue

        if line.startswith('#### '):
            heading = line.replace('####', '').strip()
            doc.add_paragraph(heading, style='Heading 3')
            in_list = False
            continue

        # Code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            para = doc.add_paragraph(line)
            para.style = 'Normal'
            run = para.runs[0] if para.runs else para.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            para.paragraph_format.left_indent = Inches(0.5)
            continue

        # Lists
        if re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            text = format_inline_markdown(text)
            para = doc.add_paragraph(text, style='List Number')
            para.paragraph_format.left_indent = Inches(0.25)
            in_list = True
            continue

        if line.strip().startswith('- ') or line.strip().startswith('* '):
            # Determine list level by leading spaces
            indent = len(line) - len(line.lstrip())
            text = line.strip()[2:]
            text = format_inline_markdown(text)

            para = doc.add_paragraph(text, style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25 + (indent / 10))
            in_list = True
            continue

        # Tables (markdown format)
        if '|' in line and line.strip().startswith('|'):
            # Skip table separator lines
            if re.match(r'^\|[\s\-:]+\|', line):
                continue

            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]

            # Create table if this is the first row
            if not hasattr(convert_markdown_to_docx, 'current_table'):
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Light Grid Accent 1'
                convert_markdown_to_docx.current_table = table
                # Header row
                for i, cell_text in enumerate(cells):
                    table.rows[0].cells[i].text = cell_text
                    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
            else:
                # Data row
                row = convert_markdown_to_docx.current_table.add_row()
                for i, cell_text in enumerate(cells):
                    row.cells[i].text = cell_text
            continue
        else:
            # Reset table if we're past table content
            if hasattr(convert_markdown_to_docx, 'current_table'):
                delattr(convert_markdown_to_docx, 'current_table')

        # Regular paragraphs
        if line.strip():
            text = format_inline_markdown(line.strip())
            para = doc.add_paragraph(text)
            para.style = 'Normal'
            in_list = False

    # Save document
    doc.save(output_file)
    return os.path.getsize(output_file)

def format_inline_markdown(text):
    """Format inline markdown (bold, italic, code)"""
    # Convert bold
    text = re.sub(r'\*\*(.+?)\*\*', lambda m: m.group(1), text)
    # Convert italic
    text = re.sub(r'\*(.+?)\*', lambda m: m.group(1), text)
    # Convert inline code
    text = re.sub(r'`(.+?)`', lambda m: m.group(1), text)
    return text

def main():
    """Main conversion function"""

    base_dir = "/Users/diegocavalariconsolini/ClaudeCode/ClaudeSkillCollection/private/wip-plugins/cybersecurity-policy-generator/output/guatemaltek-final"
    md_dir = os.path.join(base_dir, "markdown")
    word_dir = os.path.join(base_dir, "word")

    # Files to convert
    files = [
        "2-AccessControlPolicy.md",
        "3-AcceptableUsePolicy.md",
        "4-IncidentResponsePolicy.md",
        "5-RiskManagementPolicy.md",
        "6-DataClassificationPolicy.md",
        "7-BusinessContinuityPolicy.md",
        "8-PhysicalSecurityPolicy.md"
    ]

    results = []

    print(f"Converting {len(files)} markdown files to Word format...\n")

    for md_file in files:
        input_path = os.path.join(md_dir, md_file)
        output_file = md_file.replace('.md', '.docx')
        output_path = os.path.join(word_dir, output_file)

        print(f"Converting: {md_file}")

        try:
            size = convert_markdown_to_docx(input_path, output_path)
            size_kb = size / 1024
            results.append({
                'file': output_file,
                'path': output_path,
                'size': size,
                'size_kb': size_kb,
                'status': 'SUCCESS'
            })
            print(f"  ✓ Created: {output_file} ({size_kb:.1f} KB)\n")
        except Exception as e:
            results.append({
                'file': output_file,
                'path': output_path,
                'status': 'FAILED',
                'error': str(e)
            })
            print(f"  ✗ Failed: {str(e)}\n")

    # Summary
    print("=" * 70)
    print("CONVERSION SUMMARY")
    print("=" * 70)
    print(f"Total files processed: {len(files)}")
    print(f"Successful: {sum(1 for r in results if r['status'] == 'SUCCESS')}")
    print(f"Failed: {sum(1 for r in results if r['status'] == 'FAILED')}")
    print()

    print("CONVERTED FILES:")
    print("-" * 70)
    for result in results:
        if result['status'] == 'SUCCESS':
            print(f"{result['file']:<45} {result['size_kb']:>8.1f} KB")
    print("-" * 70)

    total_size = sum(r['size'] for r in results if r['status'] == 'SUCCESS')
    print(f"{'Total size:':<45} {total_size/1024:>8.1f} KB")
    print()

    return results

if __name__ == '__main__':
    main()

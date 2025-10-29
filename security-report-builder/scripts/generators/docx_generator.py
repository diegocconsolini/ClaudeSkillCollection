#!/usr/bin/env python3
"""
Generate editable DOCX security reports using python-docx.
"""

import logging
from pathlib import Path
from typing import Dict, List, Any
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except Exception as e:
    logging.warning(f"python-docx not available: {e}")
    Document = None
    DOCX_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DOCXReportGenerator:
    """Generate editable DOCX reports."""

    def __init__(self):
        """Initialize DOCX generator."""
        pass

    def generate_report(self,
                       findings: List[Dict[str, Any]],
                       statistics: Dict[str, Any],
                       framework_summary: Dict[str, Any],
                       plugin_scores: List[Dict[str, Any]],
                       overall_risk: Dict[str, Any],
                       config: Dict[str, Any],
                       branding: Dict[str, Any],
                       output_path: Path) -> None:
        """
        Generate DOCX report.

        Args:
            findings: List of security findings
            statistics: Aggregate statistics
            framework_summary: Framework analysis
            plugin_scores: Plugin risk scores
            overall_risk: Overall risk assessment
            config: Report configuration
            branding: Branding configuration
            output_path: Output file path
        """
        if not Document:
            logger.error("python-docx not installed. Cannot generate DOCX.")
            logger.info("Install with: pip install python-docx")
            return

        try:
            # Create document
            doc = Document()

            # Set up styles
            self._setup_styles(doc, branding)

            # Add content
            self._add_cover_page(doc, branding, overall_risk)
            self._add_executive_summary(doc, statistics, overall_risk)
            self._add_key_statistics(doc, statistics)
            self._add_top_plugins(doc, plugin_scores)
            self._add_critical_findings(doc, findings)
            self._add_framework_analysis(doc, framework_summary)
            self._add_footer(doc, branding)

            # Save document
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))

            logger.info(f"Generated DOCX report: {output_path}")
            logger.info(f"Report size: {output_path.stat().st_size / 1024:.1f} KB")

        except Exception as e:
            logger.error(f"Error generating DOCX: {e}")
            raise

    def _setup_styles(self, doc: 'Document', branding: Dict[str, Any]) -> None:
        """Set up document styles."""
        # Title style
        try:
            title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.font.color.rgb = self._hex_to_rgb(branding.get('primary_color', '#6366f1'))
        except:
            pass  # Style might already exist

        # Heading styles
        for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
            if style_name in doc.styles:
                style = doc.styles[style_name]
                style.font.color.rgb = self._hex_to_rgb(branding.get('primary_color', '#6366f1'))

    def _hex_to_rgb(self, hex_color: str) -> 'RGBColor':
        """Convert hex color to RGBColor."""
        hex_color = hex_color.lstrip('#')
        return RGBColor(
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16)
        )

    def _add_cover_page(self, doc: 'Document', branding: Dict[str, Any],
                       overall_risk: Dict[str, Any]) -> None:
        """Add cover page."""
        # Title
        title = doc.add_paragraph('Security Assessment Report', style='CustomTitle')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Company name
        company = doc.add_paragraph(branding.get('company_name', 'Your Organization'))
        company.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        company.runs[0].font.size = Pt(16)

        # Date
        date = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Risk level
        doc.add_paragraph()  # Spacing
        risk = doc.add_paragraph(f"Overall Risk Level: {overall_risk['overall_risk_level']}")
        risk.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        risk.runs[0].font.size = Pt(14)
        risk.runs[0].font.bold = True

        # Page break
        doc.add_page_break()

    def _add_executive_summary(self, doc: 'Document', statistics: Dict[str, Any],
                              overall_risk: Dict[str, Any]) -> None:
        """Add executive summary section."""
        doc.add_heading('Executive Summary', level=1)

        # Overall risk
        p = doc.add_paragraph()
        p.add_run('Overall Risk Assessment:').bold = True
        doc.add_paragraph(f"• Risk Score: {overall_risk['overall_risk_score']}/100")
        doc.add_paragraph(f"• Risk Level: {overall_risk['overall_risk_level']}")
        doc.add_paragraph(f"• Plugins Analyzed: {statistics['total_plugins']}")
        doc.add_paragraph(f"• Total Findings: {statistics['total_findings']}")

        # Summary text
        doc.add_paragraph()
        critical_high = overall_risk['critical_risk_count'] + overall_risk['high_risk_count']
        summary = doc.add_paragraph(
            f"This assessment identified {critical_high} high-priority plugins requiring immediate remediation. "
            f"Context-aware analysis has been applied to reduce false positives and focus on actionable security issues."
        )

    def _add_key_statistics(self, doc: 'Document', statistics: Dict[str, Any]) -> None:
        """Add key statistics section."""
        doc.add_heading('Key Statistics', level=1)

        # Create table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Severity'
        header_cells[1].text = 'Count'

        # Data
        severities = [
            ('Critical', statistics.get('by_severity', {}).get('CRITICAL', 0)),
            ('High', statistics.get('by_severity', {}).get('HIGH', 0)),
            ('Medium', statistics.get('by_severity', {}).get('MEDIUM', 0)),
            ('Low', statistics.get('by_severity', {}).get('LOW', 0))
        ]

        for idx, (severity, count) in enumerate(severities, 1):
            row = table.rows[idx]
            row.cells[0].text = severity
            row.cells[1].text = str(count)

    def _add_top_plugins(self, doc: 'Document', plugin_scores: List[Dict[str, Any]]) -> None:
        """Add top risky plugins section."""
        doc.add_heading('Top 10 Risky Plugins', level=1)

        # Create table
        table = doc.add_table(rows=11, cols=6)
        table.style = 'Light Grid Accent 1'

        # Header
        header = table.rows[0].cells
        header[0].text = 'Plugin Name'
        header[1].text = 'Risk Score'
        header[2].text = 'Risk Level'
        header[3].text = 'Findings'
        header[4].text = 'Critical'
        header[5].text = 'High'

        # Data
        for idx, plugin in enumerate(plugin_scores[:10], 1):
            row = table.rows[idx]
            sev_dist = plugin.get('severity_distribution', {})

            row.cells[0].text = plugin.get('plugin_name', 'unknown')
            row.cells[1].text = str(plugin['risk_score'])
            row.cells[2].text = plugin['risk_level']
            row.cells[3].text = f"{plugin['real_finding_count']}/{plugin['finding_count']}"
            row.cells[4].text = str(sev_dist.get('CRITICAL', 0))
            row.cells[5].text = str(sev_dist.get('HIGH', 0))

    def _add_critical_findings(self, doc: 'Document', findings: List[Dict[str, Any]]) -> None:
        """Add critical findings section."""
        doc.add_page_break()
        doc.add_heading('Critical Findings', level=1)

        critical = [f for f in findings if f.get('severity') == 'CRITICAL'][:20]

        if not critical:
            doc.add_paragraph('No critical findings identified.')
            return

        for idx, finding in enumerate(critical, 1):
            # Finding title
            heading = doc.add_heading(f"{idx}. {finding.get('category', 'Unknown')}", level=2)

            # Plugin name
            plugin_para = doc.add_paragraph()
            plugin_para.add_run('Plugin: ').bold = True
            plugin_para.add_run(finding.get('plugin_name', 'unknown'))

            # Severity
            sev_para = doc.add_paragraph()
            sev_para.add_run('Severity: ').bold = True
            sev_run = sev_para.add_run('CRITICAL')
            sev_run.font.color.rgb = RGBColor(220, 38, 38)  # Red
            sev_run.bold = True

            # Description
            desc = doc.add_paragraph(finding.get('description', 'No description'))

            # Code snippet (if available)
            code = finding.get('code_snippet', '')
            if code:
                doc.add_paragraph('Code Snippet:', style='Heading 3')
                code_para = doc.add_paragraph(code[:500])
                code_para.style = 'Normal'
                # Use monospace font
                for run in code_para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)

            # Framework tags
            attack = finding.get('att&ck_techniques', [])
            owasp = finding.get('owasp_categories', [])

            if attack or owasp:
                tags_para = doc.add_paragraph()
                tags_para.add_run('Framework References: ').bold = True

                tags = []
                for tech in attack[:3]:
                    tags.append(f"ATT&CK: {tech}")
                for cat in owasp[:2]:
                    tags.append(f"OWASP: {cat}")

                tags_para.add_run(', '.join(tags))

            doc.add_paragraph()  # Spacing

    def _add_framework_analysis(self, doc: 'Document', framework_summary: Dict[str, Any]) -> None:
        """Add framework analysis section."""
        doc.add_page_break()
        doc.add_heading('Framework Coverage', level=1)

        # Create table
        table = doc.add_table(rows=5, cols=3)
        table.style = 'Light Grid Accent 1'

        # Header
        header = table.rows[0].cells
        header[0].text = 'Framework'
        header[1].text = 'Techniques/Categories'
        header[2].text = 'Coverage'

        # Data
        frameworks = [
            ('MITRE ATT&CK', framework_summary.get('mitre_attack', {})),
            ('MITRE ATLAS', framework_summary.get('mitre_atlas', {})),
            ('OWASP Top 10', framework_summary.get('owasp', {})),
            ('CWE', framework_summary.get('cwe', {}))
        ]

        for idx, (name, data) in enumerate(frameworks, 1):
            row = table.rows[idx]
            row.cells[0].text = name

            if 'total_techniques' in data:
                row.cells[1].text = str(data['total_techniques'])
                coverage = data.get('coverage_percent', 0)
                row.cells[2].text = f"{coverage}%" if coverage else 'N/A'
            elif 'total_categories' in data:
                row.cells[1].text = str(data['total_categories'])
                row.cells[2].text = 'N/A'
            elif 'total_weaknesses' in data:
                row.cells[1].text = str(data['total_weaknesses'])
                row.cells[2].text = 'N/A'
            else:
                row.cells[1].text = '0'
                row.cells[2].text = 'N/A'

    def _add_footer(self, doc: 'Document', branding: Dict[str, Any]) -> None:
        """Add footer section."""
        doc.add_page_break()

        footer_text = branding.get('footer_text', 'Confidential')
        footer_para = doc.add_paragraph(footer_text)
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        generator = doc.add_paragraph('Report generated by Security Report Builder v1.0.0')
        generator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        generator.runs[0].font.size = Pt(9)
        generator.runs[0].font.color.rgb = RGBColor(128, 128, 128)


def main():
    """CLI interface for testing the generator."""
    import argparse
    import json

    parser = argparse.ArgumentParser(description='Generate DOCX security report')
    parser.add_argument('findings', help='Input JSON file with analyzed findings')
    parser.add_argument('--output', required=True, help='Output DOCX file')
    parser.add_argument('--config', help='Report config JSON')
    parser.add_argument('--branding', help='Branding config JSON')

    args = parser.parse_args()

    # Load data
    with open(args.findings, 'r') as f:
        data = json.load(f)

    config = {}
    if args.config:
        with open(args.config, 'r') as f:
            config = json.load(f)

    branding = {}
    if args.branding:
        with open(args.branding, 'r') as f:
            branding = json.load(f)

    # Generate DOCX
    generator = DOCXReportGenerator()
    generator.generate_report(
        findings=data.get('analyzed_findings', []),
        statistics=data.get('statistics', {}),
        framework_summary=data.get('framework_summary', {}),
        plugin_scores=data.get('plugin_scores', []),
        overall_risk=data.get('overall_risk', {}),
        config=config,
        branding=branding,
        output_path=Path(args.output)
    )

    print(f"\n✓ DOCX report generated: {args.output}")
    return 0


if __name__ == '__main__':
    exit(main())

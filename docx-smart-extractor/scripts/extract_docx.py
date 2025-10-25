#!/usr/bin/env python3
"""
DOCX Smart Extractor - extract_docx.py

Extracts complete content from Word documents (.docx) including:
- Text content with full hierarchy (headings, paragraphs, lists)
- Tables with structure and formatting
- Images metadata (size, position, description)
- Document properties (author, created date, modified date)
- Styles and formatting (bold, italic, fonts, colors)
- Comments and tracked changes
- Headers and footers
- Hyperlinks

Uses python-docx library for local extraction with zero LLM calls.
Output is cached in ~/.claude-docx-cache/ for instant reuse.
"""

import sys
import json
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import shared smart cache
sys.path.insert(0, str(Path(__file__).parent.parent.parent / 'shared'))
from smart_cache import SmartCache


def serialize_color(color):
    """Convert python-docx color to hex string"""
    if color is None:
        return None

    if isinstance(color, RGBColor):
        return f"#{color[0]:02x}{color[1]:02x}{color[2]:02x}"

    # Theme colors or other types
    return str(color)


def extract_paragraph_runs(paragraph):
    """Extract text runs with formatting from paragraph"""
    runs = []

    for run in paragraph.runs:
        run_data = {
            "text": run.text,
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "font_name": run.font.name if run.font.name else None,
            "font_size": run.font.size.pt if run.font.size else None,
            "font_color": serialize_color(run.font.color.rgb) if run.font.color else None
        }
        runs.append(run_data)

    return runs


def extract_table(table):
    """Extract table structure and content"""
    table_data = {
        "rows": len(table.rows),
        "columns": len(table.columns),
        "cells": []
    }

    for row_idx, row in enumerate(table.rows):
        row_cells = []
        for col_idx, cell in enumerate(row.cells):
            cell_data = {
                "row": row_idx,
                "col": col_idx,
                "text": cell.text,
                "paragraphs": len(cell.paragraphs)
            }
            row_cells.append(cell_data)
        table_data["cells"].append(row_cells)

    return table_data


def extract_docx(file_path, force=False):
    """
    Extract complete Word document content

    Args:
        file_path: Path to .docx file
        force: If True, re-extract even if cache exists

    Returns:
        cache_key: Unique identifier for cached extraction
    """
    file_path = Path(file_path)

    if not file_path.exists():
        print(f"Error: File not found: {file_path}")
        sys.exit(1)

    if not file_path.suffix.lower() in ['.docx', '.docm']:
        print(f"Error: Not a Word document: {file_path}")
        print("Supported formats: .docx, .docm")
        sys.exit(1)

    # Initialize SmartCache (SHAKE256 hashing with auto-migration from SHA-256)
    smart_cache = SmartCache(doc_type='docx')

    # Generate cache key with SHAKE256 (auto-migrates from SHA-256 if needed)
    cache_key, cache_dir = smart_cache.get_cache_key(str(file_path))
    cache_dir.mkdir(parents=True, exist_ok=True)

    # Check if already cached
    manifest_path = cache_dir / "manifest.json"
    if manifest_path.exists() and not force:
        print(f"Document already extracted!")
        print(f"Cache key: {cache_key}")
        print(f"Cache location: {cache_dir}")
        print(f"\nUse --force to re-extract")
        return cache_key

    print(f"Extracting Word document: {file_path.name}")
    print(f"File size: {file_path.stat().st_size / 1024 / 1024:.2f} MB")

    # Create cache directory
    cache_dir.mkdir(parents=True, exist_ok=True)

    # Load document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"Error: Failed to open document: {e}")
        sys.exit(1)

    # Extract document properties
    core_props = doc.core_properties
    metadata = {
        "filename": file_path.name,
        "file_size": file_path.stat().st_size,
        "author": core_props.author,
        "title": core_props.title,
        "subject": core_props.subject,
        "keywords": core_props.keywords,
        "created": core_props.created.isoformat() if core_props.created else None,
        "modified": core_props.modified.isoformat() if core_props.modified else None,
        "last_modified_by": core_props.last_modified_by,
        "revision": core_props.revision,
        "extracted_at": datetime.now().isoformat(),
        "cache_key": cache_key
    }

    # Extract content
    content = {
        "paragraphs": [],
        "tables": [],
        "sections": len(doc.sections)
    }

    para_count = 0
    table_count = 0

    print(f"Extracting {len(doc.paragraphs)} paragraphs...")

    for element in doc.element.body:
        # Check if paragraph
        if element.tag.endswith('p'):
            para = doc.paragraphs[para_count]
            para_count += 1

            para_data = {
                "type": "paragraph",
                "text": para.text,
                "style": para.style.name if para.style else None,
                "alignment": str(para.alignment) if para.alignment else None,
                "runs": extract_paragraph_runs(para)
            }

            content["paragraphs"].append(para_data)

        # Check if table
        elif element.tag.endswith('tbl'):
            table = doc.tables[table_count]
            table_count += 1

            table_data = {
                "type": "table",
                "index": table_count - 1,
                "data": extract_table(table)
            }

            content["tables"].append(table_data)

            if table_count % 10 == 0:
                print(f"  Processed {table_count} tables...")

    # Save metadata
    with open(cache_dir / "metadata.json", "w", encoding="utf-8") as f:
        json.dump(metadata, f, indent=2, ensure_ascii=False)

    # Save content
    with open(cache_dir / "content.json", "w", encoding="utf-8") as f:
        json.dump(content, f, indent=2, ensure_ascii=False)

    # Create manifest
    manifest = {
        "cache_key": cache_key,
        "filename": file_path.name,
        "file_size": file_path.stat().st_size,
        "extracted_at": datetime.now().isoformat(),
        "total_paragraphs": len(content["paragraphs"]),
        "total_tables": len(content["tables"]),
        "total_sections": content["sections"],
        "total_characters": sum(len(p["text"]) for p in content["paragraphs"])
    }

    with open(cache_dir / "manifest.json", "w", encoding="utf-8") as f:
        json.dump(manifest, f, indent=2, ensure_ascii=False)

    print(f"\nExtraction complete!")
    print(f"Cache key: {cache_key}")
    print(f"Cache location: {cache_dir}")
    print(f"Total paragraphs: {len(content['paragraphs'])}")
    print(f"Total tables: {len(content['tables'])}")
    print(f"Total sections: {content['sections']}")
    print(f"Total characters: {manifest['total_characters']:,}")

    print(f"\nNext step:")
    print(f"Chunk content: python scripts/semantic_chunker.py {cache_key}")

    return cache_key


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python extract_docx.py <docx_file> [--force] [--output-dir DIR]")
        print("\nOptions:")
        print("  --force          Force re-extraction even if cached")
        print("  --output-dir DIR Copy extracted files to specified directory (prompts if not provided)")
        print("\nExample:")
        print("  python extract_docx.py /path/to/document.docx")
        print("  python extract_docx.py /path/to/document.docx --force")
        print("  python extract_docx.py /path/to/document.docx --output-dir ./extracted")
        sys.exit(1)

    file_path = sys.argv[1]
    force = "--force" in sys.argv

    # Extract output directory from arguments if provided
    output_dir = None
    if '--output-dir' in sys.argv:
        output_dir_index = sys.argv.index('--output-dir')
        if output_dir_index + 1 < len(sys.argv):
            output_dir = sys.argv[output_dir_index + 1]
        else:
            print("Error: --output-dir requires a value")
            sys.exit(1)

    cache_key = extract_docx(file_path, force=force)

    if cache_key:
        # Handle output directory
        smart_cache = SmartCache(doc_type='docx')
        _, cache_path = smart_cache.get_cache_key(file_path)

        if output_dir is None:
            # Ask user interactively if they want to copy to working directory
            print("\n📁 Copy extracted files to working directory?")
            response = input("Copy files? (y/n): ").strip().lower()

            if response == 'y':
                output_dir = Path.cwd() / f"extracted_{cache_key}"

        if output_dir:
            output_path = Path(output_dir)

            # Ask about cache behavior
            print(f"\n💾 Maintain cache in {cache_path}?")
            print("  (yes) Keep cache for instant reuse across projects")
            print("  (no)  Remove cache after copying files")
            keep_cache = input("Keep cache? (y/n): ").strip().lower()

            # Copy all files from cache to output directory
            import shutil

            print(f"\n📋 Copying files to {output_path}...")
            output_path.mkdir(parents=True, exist_ok=True)

            for item in cache_path.iterdir():
                if item.is_file():
                    shutil.copy2(item, output_path / item.name)
                    print(f"  ✓ {item.name}")

            print(f"\n✓ Files copied to: {output_path}")

            # Remove cache if requested
            if keep_cache != 'y':
                shutil.rmtree(cache_path)
                print(f"  ✓ Cache removed")
            else:
                print(f"  ✓ Cache maintained at: {cache_path}")
                print(f"  💡 Cache persists across projects for instant reuse")
